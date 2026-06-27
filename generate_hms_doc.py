"""
generate_hms_doc.py
Generates: Working_Functionalities_of_HMS.docx
with embedded flowchart images drawn via matplotlib.
"""

import os
import sys
import io
import textwrap
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import matplotlib.patheffects as pe
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.font_manager as fm

# ─── Color Palette ────────────────────────────────────────────────────────────
PRIMARY   = '#6366f1'
TEAL      = '#0d9488'
GREEN     = '#10b981'
AMBER     = '#f59e0b'
RED       = '#ef4444'
BLUE      = '#0ea5e9'
PURPLE    = '#8b5cf6'
GRAY      = '#64748b'
LIGHTGRAY = '#f1f5f9'
WHITE     = '#ffffff'
DARK      = '#1e293b'

# ─── Helper: save figure to bytes ─────────────────────────────────────────────
def fig_to_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor=fig.get_facecolor())
    buf.seek(0)
    plt.close(fig)
    return buf

# ─── Helper: draw a rounded box ───────────────────────────────────────────────
def draw_box(ax, x, y, w, h, text, bg=PRIMARY, fg=WHITE,
             fontsize=8, radius=0.04, bold=False):
    box = FancyBboxPatch((x - w/2, y - h/2), w, h,
                         boxstyle=f"round,pad=0.01,rounding_size={radius}",
                         linewidth=1.2, edgecolor=bg,
                         facecolor=bg, zorder=3)
    ax.add_patch(box)
    weight = 'bold' if bold else 'normal'
    ax.text(x, y, text, ha='center', va='center',
            fontsize=fontsize, color=fg, fontweight=weight,
            zorder=4, wrap=True,
            multialignment='center')

def draw_diamond(ax, x, y, w, h, text, bg=AMBER, fg=WHITE, fontsize=7.5):
    dx, dy = w/2, h/2
    pts = [[x, y+dy],[x+dx, y],[x, y-dy],[x-dx, y]]
    diamond = plt.Polygon(pts, closed=True, facecolor=bg,
                          edgecolor=bg, linewidth=1.2, zorder=3)
    ax.add_patch(diamond)
    ax.text(x, y, text, ha='center', va='center',
            fontsize=fontsize, color=fg, fontweight='bold', zorder=4,
            multialignment='center')

def arrow(ax, x1, y1, x2, y2, label='', color=GRAY):
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle='->', color=color,
                                lw=1.4, connectionstyle='arc3,rad=0'))
    if label:
        mx, my = (x1+x2)/2, (y1+y2)/2
        ax.text(mx+0.03, my, label, fontsize=6.5, color=color,
                ha='left', va='center', style='italic')

def new_fig(w=9, h=5.5, bg=LIGHTGRAY):
    fig, ax = plt.subplots(figsize=(w, h))
    fig.patch.set_facecolor(bg)
    ax.set_facecolor(bg)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis('off')
    return fig, ax

def add_title(ax, text, fontsize=11):
    ax.text(0.5, 0.97, text, ha='center', va='top',
            fontsize=fontsize, fontweight='bold', color=DARK,
            transform=ax.transAxes)

# ══════════════════════════════════════════════════════════════════════════════
# FLOWCHART 1 — System Architecture
# ══════════════════════════════════════════════════════════════════════════════
def chart_system_arch():
    fig, ax = new_fig(9, 4.5)
    add_title(ax, 'SwasthyaSync — System Architecture')

    nodes = [
        (0.10, 0.50, 0.14, 0.12, 'Browser\n/ Client',            BLUE,    WHITE),
        (0.30, 0.50, 0.14, 0.12, 'React 18\nFrontend (Vite)',    PRIMARY, WHITE),
        (0.50, 0.50, 0.14, 0.12, 'Route Guards\n& AuthContext',  PURPLE,  WHITE),
        (0.70, 0.75, 0.14, 0.12, 'Express.js\nREST API',         TEAL,    WHITE),
        (0.70, 0.25, 0.14, 0.12, 'Login Page\n/login',           RED,     WHITE),
        (0.88, 0.85, 0.14, 0.12, 'PostgreSQL\nDatabase',         AMBER,   WHITE),
        (0.88, 0.65, 0.14, 0.12, 'File Storage\n/uploads/',      GREEN,   WHITE),
        (0.88, 0.45, 0.14, 0.12, 'Email Service\nNodemailer',    BLUE,    WHITE),
    ]
    for (x, y, w, h, txt, bg, fg) in nodes:
        draw_box(ax, x, y, w, h, txt, bg=bg, fg=fg, fontsize=7.5)

    arrows = [
        (0.17, 0.50, 0.23, 0.50, ''),
        (0.37, 0.50, 0.43, 0.50, ''),
        (0.57, 0.55, 0.63, 0.73, 'Authenticated'),
        (0.57, 0.45, 0.63, 0.27, 'Unauth → Login'),
        (0.77, 0.75, 0.81, 0.85, ''),
        (0.77, 0.75, 0.81, 0.65, ''),
        (0.77, 0.75, 0.81, 0.45, ''),
    ]
    for (x1,y1,x2,y2,lbl) in arrows:
        arrow(ax, x1, y1, x2, y2, lbl)
    return fig_to_bytes(fig)

# ══════════════════════════════════════════════════════════════════════════════
# FLOWCHART 2 — Authentication Flow
# ══════════════════════════════════════════════════════════════════════════════
def chart_auth():
    fig, ax = new_fig(8, 5.5)
    add_title(ax, 'Authentication & Access Control Flow')

    draw_box(ax, 0.15, 0.82, 0.20, 0.10, 'User visits\n/app/*', BLUE, WHITE, 7.5)
    draw_diamond(ax, 0.15, 0.62, 0.22, 0.10, 'Token in\nlocalStorage?', AMBER, WHITE, 7)
    draw_box(ax, 0.15, 0.40, 0.20, 0.10, 'Login Form\nEmail + Password', RED, WHITE, 7.5)
    draw_box(ax, 0.15, 0.20, 0.20, 0.10, 'POST /api/auth/login', GRAY, WHITE, 7.5)

    draw_box(ax, 0.50, 0.62, 0.20, 0.10, 'Verify JWT\nwith server', PRIMARY, WHITE, 7.5)
    draw_diamond(ax, 0.50, 0.42, 0.22, 0.10, 'Valid\nToken?', AMBER, WHITE, 7)
    draw_diamond(ax, 0.50, 0.22, 0.22, 0.10, 'User\nRole?', PURPLE, WHITE, 7)
    draw_box(ax, 0.80, 0.32, 0.18, 0.10, 'ADMIN\nFull Access\n40+ modules', GREEN, WHITE, 7)
    draw_box(ax, 0.80, 0.15, 0.18, 0.10, 'STAFF\nStandard Access\n(no admin)', TEAL, WHITE, 7)

    arrow(ax, 0.15, 0.77, 0.15, 0.67)
    arrow(ax, 0.15, 0.57, 0.15, 0.45, 'No Token')
    arrow(ax, 0.15, 0.35, 0.15, 0.25)
    arrow(ax, 0.26, 0.62, 0.40, 0.62, 'Has Token')
    arrow(ax, 0.50, 0.57, 0.50, 0.47)
    arrow(ax, 0.50, 0.37, 0.50, 0.27)
    arrow(ax, 0.38, 0.42, 0.15, 0.42, 'Invalid →')
    arrow(ax, 0.61, 0.22, 0.71, 0.30, 'Admin')
    arrow(ax, 0.61, 0.20, 0.71, 0.17, 'Staff')
    return fig_to_bytes(fig)

# ══════════════════════════════════════════════════════════════════════════════
# FLOWCHART 3 — Patient Lifecycle
# ══════════════════════════════════════════════════════════════════════════════
def chart_patient_lifecycle():
    fig, ax = new_fig(10, 5.5)
    add_title(ax, 'Patient Lifecycle — End-to-End Workflow')

    draw_box(ax, 0.10, 0.75, 0.14, 0.10, 'Walk-in /\nReferred Patient', PRIMARY, WHITE, 7)
    draw_box(ax, 0.10, 0.55, 0.14, 0.10, 'Reception\nRegistration\n(UHID generated)', TEAL, WHITE, 6.5)
    draw_diamond(ax, 0.10, 0.35, 0.16, 0.10, 'Admission\nType?', AMBER, WHITE, 6.5)

    draw_box(ax, 0.30, 0.65, 0.13, 0.09, 'OPD Queue\nToken', GREEN, WHITE, 7)
    draw_box(ax, 0.30, 0.45, 0.13, 0.09, 'IPD Bed\nAssignment', BLUE, WHITE, 7)
    draw_box(ax, 0.30, 0.25, 0.13, 0.09, 'Emergency\nTriage', RED, WHITE, 7)

    draw_box(ax, 0.50, 0.55, 0.13, 0.09, 'Consultation\n& Clinical Notes', PRIMARY, WHITE, 7)
    draw_box(ax, 0.50, 0.35, 0.13, 0.09, 'Ward / Nursing\nStation', PURPLE, WHITE, 7)

    draw_box(ax, 0.68, 0.65, 0.13, 0.09, 'Lab Orders\n(LIS)', TEAL, WHITE, 7)
    draw_box(ax, 0.68, 0.45, 0.13, 0.09, 'Radiology\n(RIS)', BLUE, WHITE, 7)
    draw_box(ax, 0.68, 0.25, 0.13, 0.09, 'Pharmacy\nDispensing', GREEN, WHITE, 7)

    draw_box(ax, 0.86, 0.55, 0.13, 0.09, 'Discharge\nSummary PDF', AMBER, WHITE, 7)
    draw_box(ax, 0.86, 0.35, 0.13, 0.09, 'Final Billing\n& Settlement', RED, WHITE, 7)
    draw_box(ax, 0.86, 0.15, 0.13, 0.09, 'Bed Released\nAudit Logged', GRAY, WHITE, 7)

    arrow(ax, 0.10, 0.70, 0.10, 0.60)
    arrow(ax, 0.10, 0.50, 0.10, 0.40)
    arrow(ax, 0.18, 0.37, 0.24, 0.65, 'OPD')
    arrow(ax, 0.18, 0.35, 0.24, 0.45, 'IPD')
    arrow(ax, 0.18, 0.33, 0.24, 0.25, 'Emergency')
    arrow(ax, 0.37, 0.65, 0.44, 0.57)
    arrow(ax, 0.37, 0.45, 0.44, 0.38)
    arrow(ax, 0.37, 0.25, 0.44, 0.37)
    arrow(ax, 0.57, 0.55, 0.62, 0.65)
    arrow(ax, 0.57, 0.52, 0.62, 0.45)
    arrow(ax, 0.57, 0.38, 0.62, 0.27)
    arrow(ax, 0.75, 0.65, 0.80, 0.57)
    arrow(ax, 0.75, 0.45, 0.80, 0.52)
    arrow(ax, 0.75, 0.25, 0.80, 0.37)
    arrow(ax, 0.93, 0.50, 0.93, 0.40)
    arrow(ax, 0.93, 0.30, 0.93, 0.20)
    return fig_to_bytes(fig)

# ══════════════════════════════════════════════════════════════════════════════
# FLOWCHART 4 — Patient Registration 4-Step Wizard
# ══════════════════════════════════════════════════════════════════════════════
def chart_registration_wizard():
    fig, ax = new_fig(10, 4)
    add_title(ax, 'New Patient Registration — 4-Step Wizard')

    steps = [
        (0.12, 0.55, 'Step 1\nPersonal Details\n• Name (req)\n• Age / DOB\n• Gender / Blood Group\n• Phone (req)\n• Address / Aadhaar / ABHA\n• Guardian contact', PRIMARY),
        (0.37, 0.55, 'Step 2\nClinical Info\n• Department (req)\n• Attending Doctor\n• Status / MLC Type\n• Patient Category\n• Chief Complaint', TEAL),
        (0.62, 0.55, 'Step 3\nPayment / Insurance\n• Payment Type\n• Insurance Company\n• TPA Name\n• Policy No / Validity', PURPLE),
        (0.87, 0.55, 'Step 4\nConsent & Submit\n• Summary preview\n• Consent checkbox (req)\n• Estimate provided\n• UHID auto-generated', GREEN),
    ]
    for (x, y, txt, bg) in steps:
        draw_box(ax, x, y, 0.22, 0.60, txt, bg=bg, fg=WHITE, fontsize=7, radius=0.03)

    for i in range(3):
        x1 = 0.12 + i*0.25 + 0.11
        x2 = x1 + 0.03
        arrow(ax, x1, 0.55, x2, 0.55, 'Validate →')

    draw_box(ax, 0.50, 0.12, 0.30, 0.10, 'Patient record created & saved\nUHID issued, redirected to patient list', GREEN, WHITE, 8, bold=True)
    arrow(ax, 0.87, 0.25, 0.60, 0.17)

    ax.text(0.12, 0.14, 'Validate each step\nbefore proceeding →',
            ha='center', va='center', fontsize=7, color=GRAY, style='italic')
    return fig_to_bytes(fig)

# ══════════════════════════════════════════════════════════════════════════════
# FLOWCHART 5 — IPD Bed Management
# ══════════════════════════════════════════════════════════════════════════════
def chart_ipd():
    fig, ax = new_fig(10, 5.5)
    add_title(ax, 'IPD Module — Bed Management & Discharge Workflow')

    draw_box(ax, 0.10, 0.82, 0.16, 0.10, 'Patient Admitted\n(IPD type)', PRIMARY, WHITE, 7.5)
    draw_diamond(ax, 0.10, 0.62, 0.18, 0.10, 'Bed\nAvailable?', AMBER, WHITE, 7)
    draw_box(ax, 0.10, 0.42, 0.16, 0.10, 'Wait / Ward\nTransfer', RED, WHITE, 7.5)
    draw_box(ax, 0.32, 0.82, 0.16, 0.10, 'Step 1: Select Bed\nGrid of available beds', TEAL, WHITE, 7)
    draw_box(ax, 0.32, 0.62, 0.16, 0.10, 'Step 2: Patient Details\nDoctor · Diagnosis\nCritical flag', TEAL, WHITE, 7)
    draw_box(ax, 0.32, 0.42, 0.16, 0.10, 'Bed → OCCUPIED\nAdmission recorded', GREEN, WHITE, 7.5)

    draw_box(ax, 0.57, 0.82, 0.16, 0.10, 'Patient Info Tab\nWard · Doctor\nDiagnosis', BLUE, WHITE, 7)
    draw_box(ax, 0.57, 0.62, 0.16, 0.10, 'Forms Tab\nAssign consent\nforms via FormViewer', PURPLE, WHITE, 7)
    draw_box(ax, 0.57, 0.42, 0.16, 0.10, 'Discharge Tab\nCreate discharge\nsummary PDF', AMBER, WHITE, 7)

    draw_box(ax, 0.82, 0.62, 0.16, 0.10, 'Bed → AVAILABLE\nDischarge recorded', GREEN, WHITE, 7.5)
    draw_box(ax, 0.82, 0.42, 0.16, 0.10, 'IPD Billing\n6-Step Clearance\nWorkflow', RED, WHITE, 7.5)

    # Bed state diagram
    states = [
        (0.12, 0.18, 'Available', GREEN),
        (0.35, 0.18, 'Occupied', PRIMARY),
        (0.58, 0.18, 'Maintenance', GRAY),
        (0.81, 0.18, 'Reserved', AMBER),
    ]
    for (x, y, txt, bg) in states:
        draw_box(ax, x, y, 0.17, 0.09, txt, bg=bg, fg=WHITE, fontsize=8, bold=True)
    ax.text(0.50, 0.09, 'Bed Status State Machine', ha='center', fontsize=7.5,
            color=DARK, style='italic')
    arrow(ax, 0.21, 0.18, 0.27, 0.18, 'Assign')
    arrow(ax, 0.44, 0.18, 0.50, 0.18, 'Maintenance')
    arrow(ax, 0.67, 0.18, 0.73, 0.18, 'Reserve')
    arrow(ax, 0.35, 0.22, 0.12, 0.22, 'Discharge')

    arrow(ax, 0.10, 0.77, 0.10, 0.67)
    arrow(ax, 0.10, 0.57, 0.10, 0.47, 'No')
    arrow(ax, 0.19, 0.62, 0.24, 0.82, 'Yes')
    arrow(ax, 0.32, 0.77, 0.32, 0.67)
    arrow(ax, 0.32, 0.57, 0.32, 0.47)
    arrow(ax, 0.40, 0.42, 0.49, 0.82)
    arrow(ax, 0.40, 0.42, 0.49, 0.62)
    arrow(ax, 0.40, 0.42, 0.49, 0.42)
    arrow(ax, 0.65, 0.42, 0.74, 0.62)
    arrow(ax, 0.82, 0.57, 0.82, 0.47)
    return fig_to_bytes(fig)

# ══════════════════════════════════════════════════════════════════════════════
# FLOWCHART 6 — Emergency Triage
# ══════════════════════════════════════════════════════════════════════════════
def chart_emergency():
    fig, ax = new_fig(9, 5)
    add_title(ax, 'Emergency Module — Triage & MLC Workflow')

    draw_box(ax, 0.12, 0.82, 0.18, 0.10, 'Patient arrives\nEmergency', RED, WHITE, 8, bold=True)
    draw_box(ax, 0.12, 0.62, 0.18, 0.10, 'Fast-Track\nRegistration\n(minimal fields)', AMBER, WHITE, 7.5)
    draw_diamond(ax, 0.12, 0.42, 0.20, 0.11, 'Triage\nLevel?', GRAY, WHITE, 7.5)

    levels = [
        (0.38, 0.78, 'RED — Critical\nImmediate Care\nAlert flag ON', '#dc2626'),
        (0.38, 0.60, 'ORANGE — Urgent\nSeen < 15 min', '#ea580c'),
        (0.38, 0.43, 'YELLOW — Semi-Urgent\nSeen < 60 min', '#ca8a04'),
        (0.38, 0.26, 'GREEN — Non-Urgent\nStandard OPD queue', '#16a34a'),
    ]
    for (x, y, txt, bg) in levels:
        draw_box(ax, x, y, 0.20, 0.10, txt, bg=bg, fg=WHITE, fontsize=7)

    draw_diamond(ax, 0.65, 0.62, 0.18, 0.10, 'MLC\nCase?', '#7c3aed', WHITE, 7)
    draw_box(ax, 0.65, 0.42, 0.18, 0.10, 'MLC Type selection\nPolice info / FIR\ncapture', RED, WHITE, 7)
    draw_box(ax, 0.87, 0.62, 0.18, 0.10, 'Clinical Management\nICU / Ward\nor OPD discharge', TEAL, WHITE, 7)

    arrow(ax, 0.12, 0.77, 0.12, 0.67)
    arrow(ax, 0.12, 0.57, 0.12, 0.47)
    arrow(ax, 0.22, 0.78, 0.28, 0.78)
    arrow(ax, 0.22, 0.60, 0.28, 0.60)
    arrow(ax, 0.22, 0.43, 0.28, 0.43)
    arrow(ax, 0.22, 0.26, 0.28, 0.26)
    arrow(ax, 0.48, 0.78, 0.56, 0.65)
    arrow(ax, 0.48, 0.60, 0.56, 0.62)
    arrow(ax, 0.65, 0.57, 0.65, 0.47, 'Yes')
    arrow(ax, 0.74, 0.62, 0.78, 0.62, 'No')
    arrow(ax, 0.65, 0.37, 0.78, 0.60)
    return fig_to_bytes(fig)

# ══════════════════════════════════════════════════════════════════════════════
# FLOWCHART 7 — Lab Order Workflow
# ══════════════════════════════════════════════════════════════════════════════
def chart_lab():
    fig, ax = new_fig(10, 5)
    add_title(ax, 'Laboratory Information System (LIS) — Complete Workflow')

    nodes = [
        (0.08, 0.65, 0.12, 0.10, 'Doctor places\nLab Order\n(CPOE / LIS)', PRIMARY, WHITE),
        (0.24, 0.65, 0.12, 0.10, 'Order Created\nStatus: PENDING\nPriority set', GRAY, WHITE),
        (0.40, 0.65, 0.12, 0.10, 'Sample\nCollection\nStatus: IN PROGRESS', AMBER, WHITE),
        (0.56, 0.65, 0.12, 0.10, 'Parameter entry\nper test\nRef range shown', TEAL, WHITE),
        (0.72, 0.65, 0.12, 0.10, 'AUTO: HIGH/LOW\nflag detection\nResults saved', PURPLE, WHITE),
        (0.88, 0.65, 0.12, 0.10, 'Status:\nCOMPLETED\nPDF generated', GREEN, WHITE),
    ]
    for n in nodes:
        draw_box(ax, *n, fontsize=7)

    for i in range(5):
        x1 = 0.08 + i*0.16 + 0.06
        arrow(ax, x1, 0.65, x1+0.04, 0.65)

    # Delivery options
    draw_box(ax, 0.60, 0.30, 0.14, 0.10, 'Download PDF\njsPDF report', BLUE, WHITE, 7.5)
    draw_box(ax, 0.78, 0.30, 0.14, 0.10, 'Email to patient\nNodemailer SMTP', TEAL, WHITE, 7.5)
    draw_box(ax, 0.96, 0.30, 0.14, 0.10, 'Upload External\nScanned PDF', AMBER, WHITE, 7.5)

    arrow(ax, 0.88, 0.60, 0.66, 0.35, 'Download')
    arrow(ax, 0.88, 0.60, 0.82, 0.35, 'Email')
    arrow(ax, 0.94, 0.60, 0.96, 0.35, 'External')

    # PDF structure box
    pdf_items = [
        'Hospital Header (logo/image)',
        'Patient Info Box',
        'Results Table (Parameter | Result | Flag | Unit | Range)',
        'HIGH = Red  |  LOW = Amber  |  Normal = Black',
        'Hospital Footer + Page Numbers',
    ]
    ax.text(0.10, 0.42, 'Lab Report PDF Structure:', fontsize=8,
            color=DARK, fontweight='bold')
    for i, item in enumerate(pdf_items):
        color = RED if 'HIGH' in item else (AMBER if 'LOW' in item else DARK)
        ax.text(0.10, 0.35 - i*0.055, f'  • {item}', fontsize=7, color=color)
    return fig_to_bytes(fig)

# ══════════════════════════════════════════════════════════════════════════════
# FLOWCHART 8 — Clinical Forms (FormViewer)
# ══════════════════════════════════════════════════════════════════════════════
def chart_forms():
    fig, ax = new_fig(9, 5)
    add_title(ax, 'Patient Forms & Consent Management — FormViewer Workflow')

    draw_box(ax, 0.10, 0.80, 0.16, 0.10, 'Admin uploads\nPDF Template', GRAY, WHITE, 7.5)
    draw_box(ax, 0.10, 0.58, 0.16, 0.10, 'Set category:\nConsent · Nursing\nICU · OT etc.', GRAY, WHITE, 7)
    draw_box(ax, 0.32, 0.80, 0.16, 0.10, 'Doctor/Nurse\nassigns template\nto patient', PRIMARY, WHITE, 7.5)
    draw_box(ax, 0.32, 0.58, 0.16, 0.10, 'Patient form\ninstance created\non server', TEAL, WHITE, 7.5)
    draw_box(ax, 0.55, 0.80, 0.16, 0.10, 'FormViewer opens\nFull-screen PDF\nAnnotation Engine', PURPLE, WHITE, 7.5)

    # Annotation types
    ann = [
        (0.75, 0.82, 'Checkbox\nClick to toggle', GREEN),
        (0.75, 0.65, 'Text Field\nInline typing on PDF', BLUE),
        (0.75, 0.48, 'Signature\nName insertion', AMBER),
    ]
    for (x, y, txt, bg) in ann:
        draw_box(ax, x, y, 0.18, 0.10, txt, bg=bg, fg=WHITE, fontsize=7.5)

    draw_box(ax, 0.55, 0.38, 0.16, 0.10, 'Annotations saved\nto server as JSON', TEAL, WHITE, 7.5)

    # Status progression
    statuses = [
        (0.22, 0.20, 'BLANK', GRAY),
        (0.45, 0.20, 'IN PROGRESS', AMBER),
        (0.68, 0.20, 'COMPLETED', GREEN),
    ]
    ax.text(0.45, 0.12, 'Form Status Progression', ha='center', fontsize=8, color=DARK, fontweight='bold')
    for (x, y, txt, bg) in statuses:
        draw_box(ax, x, y, 0.18, 0.09, txt, bg=bg, fg=WHITE, fontsize=8, bold=True)
    arrow(ax, 0.31, 0.20, 0.36, 0.20)
    arrow(ax, 0.54, 0.20, 0.59, 0.20)

    arrow(ax, 0.10, 0.75, 0.10, 0.63)
    arrow(ax, 0.18, 0.80, 0.24, 0.80)
    arrow(ax, 0.18, 0.58, 0.24, 0.58)
    arrow(ax, 0.32, 0.75, 0.32, 0.63)
    arrow(ax, 0.40, 0.80, 0.47, 0.80)
    arrow(ax, 0.40, 0.58, 0.47, 0.80)
    arrow(ax, 0.63, 0.80, 0.66, 0.82)
    arrow(ax, 0.63, 0.78, 0.66, 0.65)
    arrow(ax, 0.63, 0.76, 0.66, 0.48)
    arrow(ax, 0.55, 0.55, 0.55, 0.43)
    return fig_to_bytes(fig)

# ══════════════════════════════════════════════════════════════════════════════
# FLOWCHART 9 — Billing & Revenue Cycle
# ══════════════════════════════════════════════════════════════════════════════
def chart_billing():
    fig, ax = new_fig(10, 5.5)
    add_title(ax, 'Billing & Revenue Cycle — Complete Workflow')

    # Charge sources
    sources = [
        (0.10, 0.85, 'OPD\nConsultation'),
        (0.10, 0.68, 'IPD\nWard Charges'),
        (0.10, 0.51, 'Laboratory\nTests'),
        (0.10, 0.34, 'Pharmacy\nDrugs'),
        (0.10, 0.17, 'Radiology\nScans'),
    ]
    for (x, y, txt) in sources:
        draw_box(ax, x, y, 0.14, 0.09, txt, bg=PRIMARY, fg=WHITE, fontsize=7.5)
        arrow(ax, 0.17, y, 0.25, 0.50)

    draw_box(ax, 0.32, 0.50, 0.14, 0.55, 'Billing\nModule\n\nConsolidate\nall charges', TEAL, WHITE, 8, bold=True)

    # Payment methods
    pmts = [
        (0.55, 0.82, 'Cash', GREEN),
        (0.55, 0.70, 'Card / UPI', BLUE),
        (0.55, 0.58, 'Insurance\n/ TPA', PURPLE),
        (0.55, 0.46, 'CGHS / ECHS\n/ ESI', AMBER),
        (0.55, 0.34, 'Cheque\n/ NEFT', GRAY),
    ]
    for (x, y, txt, bg) in pmts:
        draw_box(ax, x, y, 0.14, 0.09, txt, bg=bg, fg=WHITE, fontsize=7.5)
        arrow(ax, 0.39, 0.50, 0.48, y)

    # Status outcomes
    draw_box(ax, 0.78, 0.75, 0.14, 0.09, 'PAID', GREEN, WHITE, 9, bold=True)
    draw_box(ax, 0.78, 0.58, 0.14, 0.09, 'PENDING', AMBER, WHITE, 9, bold=True)
    draw_box(ax, 0.78, 0.41, 0.14, 0.09, 'PARTIAL', BLUE, WHITE, 9, bold=True)
    draw_box(ax, 0.78, 0.24, 0.14, 0.09, 'OVERDUE', RED, WHITE, 9, bold=True)

    for y in [0.75, 0.58, 0.41, 0.24]:
        arrow(ax, 0.62, 0.58, 0.71, y)

    # IPD 6-step discharge
    steps6 = ['Discharge\nAdvised','Clinical\nSummary','Pharmacy\nClearance',
               'Lab\nClearance','Final\nBilling','Settle\n& Vacate']
    colors6 = [PRIMARY, TEAL, GREEN, BLUE, PURPLE, RED]
    ax.text(0.50, 0.12, 'IPD Discharge — 6-Step Clearance Workflow', ha='center',
            fontsize=8, color=DARK, fontweight='bold')
    for i, (step, clr) in enumerate(zip(steps6, colors6)):
        x = 0.10 + i * 0.15
        draw_box(ax, x, 0.04, 0.13, 0.08, step, bg=clr, fg=WHITE, fontsize=6.5)
        if i < 5:
            arrow(ax, x+0.065, 0.04, x+0.085, 0.04)
    return fig_to_bytes(fig)

# ══════════════════════════════════════════════════════════════════════════════
# FLOWCHART 10 — CPOE Orders
# ══════════════════════════════════════════════════════════════════════════════
def chart_cpoe():
    fig, ax = new_fig(9, 4.5)
    add_title(ax, 'CPOE — Computerised Physician Order Entry')

    draw_box(ax, 0.12, 0.65, 0.18, 0.12, 'Doctor places order\nvia CPOE', PRIMARY, WHITE, 8, bold=True)
    draw_diamond(ax, 0.12, 0.40, 0.20, 0.12, 'Order\nType?', AMBER, WHITE, 8)

    targets = [
        (0.50, 0.82, 'Lab Order → LIS\nPending Collection\nTest catalog 50+ tests', TEAL),
        (0.50, 0.62, 'Radiology Order → RIS\nPending Scan\nX-Ray / CT / MRI / US', BLUE),
        (0.50, 0.42, 'Pharmacy Order\nDispensing Queue\nPrescription issued', GREEN),
        (0.50, 0.22, 'Dietary Order\nNutrition Team\nDiet plan assigned', AMBER),
    ]
    labels = ['Lab', 'Radiology', 'Medication', 'Diet']
    for (x, y, txt, bg), lbl in zip(targets, labels):
        draw_box(ax, x, y, 0.26, 0.12, txt, bg=bg, fg=WHITE, fontsize=7.5)
        arrow(ax, 0.22, 0.40, 0.37, y, lbl)

    priority = [
        (0.84, 0.82, 'STAT', RED),
        (0.84, 0.62, 'URGENT', AMBER),
        (0.84, 0.42, 'ROUTINE', GREEN),
    ]
    ax.text(0.84, 0.93, 'Priority', ha='center', fontsize=8, color=DARK, fontweight='bold')
    for (x, y, txt, bg) in priority:
        draw_box(ax, x, y, 0.14, 0.09, txt, bg=bg, fg=WHITE, fontsize=8, bold=True)

    arrow(ax, 0.12, 0.59, 0.12, 0.46)
    return fig_to_bytes(fig)

# ══════════════════════════════════════════════════════════════════════════════
# FLOWCHART 11 — Appointment Scheduling
# ══════════════════════════════════════════════════════════════════════════════
def chart_appointments():
    fig, ax = new_fig(9, 4.5)
    add_title(ax, 'Appointment & Scheduling Workflow')

    nodes = [
        (0.10, 0.72, 'Reception /\nPatient books', PRIMARY, 0.14, 0.10),
        (0.10, 0.50, 'Select Dept\n& Doctor', TEAL, 0.14, 0.10),
        (0.10, 0.28, 'Check Slot\nAvailability', BLUE, 0.14, 0.10),
    ]
    for (x, y, txt, bg, w, h) in nodes:
        draw_box(ax, x, y, w, h, txt, bg=bg, fg=WHITE, fontsize=7.5)
    arrow(ax, 0.10, 0.67, 0.10, 0.55)
    arrow(ax, 0.10, 0.45, 0.10, 0.33)

    draw_diamond(ax, 0.35, 0.50, 0.20, 0.12, 'Slot\nAvailable?', AMBER, WHITE, 8)
    arrow(ax, 0.17, 0.28, 0.25, 0.48)

    draw_box(ax, 0.35, 0.26, 0.18, 0.10, 'Suggest next\navailable slot', GRAY, WHITE, 7.5)
    draw_box(ax, 0.60, 0.72, 0.18, 0.10, 'Enter patient\ndetails / select existing', TEAL, WHITE, 7.5)
    draw_box(ax, 0.60, 0.52, 0.18, 0.10, 'Confirm Booking\nStatus: SCHEDULED', GREEN, WHITE, 7.5)
    draw_box(ax, 0.84, 0.72, 0.14, 0.10, 'Arrives →\nStatus: CONFIRMED\n→ Consultation', GREEN, WHITE, 7)
    draw_box(ax, 0.84, 0.52, 0.14, 0.10, 'No-show →\nStatus: CANCELLED\nSlot freed', RED, WHITE, 7)
    draw_box(ax, 0.84, 0.32, 0.14, 0.10, 'Status:\nCOMPLETED\nOPD billing', BLUE, WHITE, 7)

    arrow(ax, 0.45, 0.50, 0.51, 0.72, 'Yes')
    arrow(ax, 0.35, 0.44, 0.35, 0.31, 'No')
    arrow(ax, 0.60, 0.67, 0.60, 0.57)
    arrow(ax, 0.69, 0.72, 0.77, 0.72)
    arrow(ax, 0.69, 0.52, 0.77, 0.52)
    arrow(ax, 0.84, 0.67, 0.84, 0.57)
    arrow(ax, 0.91, 0.67, 0.91, 0.37)
    return fig_to_bytes(fig)

# ══════════════════════════════════════════════════════════════════════════════
# FLOWCHART 12 — Settings & Report Branding
# ══════════════════════════════════════════════════════════════════════════════
def chart_settings():
    fig, ax = new_fig(9, 4.5)
    add_title(ax, 'Settings Module & Report Branding Decision Flow')

    draw_box(ax, 0.12, 0.72, 0.18, 0.10, 'Generate PDF Report\n(Lab / Discharge)', PRIMARY, WHITE, 7.5, bold=True)
    draw_diamond(ax, 0.12, 0.50, 0.22, 0.12, 'report_print\n_mode?', AMBER, WHITE, 7.5)
    draw_box(ax, 0.12, 0.28, 0.20, 0.11, 'Text / Logo Mode\nHeader text + tagline\nLogo from settings', TEAL, WHITE, 7)
    draw_box(ax, 0.38, 0.28, 0.20, 0.11, 'Image Mode\nFull-page header image\nembedded in PDF', PURPLE, WHITE, 7)
    draw_box(ax, 0.65, 0.50, 0.18, 0.10, 'Report Body\nPatient info\nResults / Content', BLUE, WHITE, 7.5)
    draw_diamond(ax, 0.65, 0.30, 0.20, 0.12, 'Footer\nMode?', AMBER, WHITE, 7.5)
    draw_box(ax, 0.55, 0.12, 0.18, 0.10, 'Text Footer\nFooter text\n+ page numbers', TEAL, WHITE, 7)
    draw_box(ax, 0.80, 0.12, 0.18, 0.10, 'Image Footer\nFull-page footer\nimage embedded', PURPLE, WHITE, 7)

    arrow(ax, 0.12, 0.67, 0.12, 0.56)
    arrow(ax, 0.12, 0.44, 0.12, 0.34, 'text')
    arrow(ax, 0.23, 0.50, 0.28, 0.30, 'image')
    arrow(ax, 0.22, 0.28, 0.56, 0.50)
    arrow(ax, 0.48, 0.28, 0.56, 0.48)
    arrow(ax, 0.65, 0.45, 0.65, 0.36)
    arrow(ax, 0.55, 0.30, 0.60, 0.17, 'text')
    arrow(ax, 0.75, 0.30, 0.80, 0.17, 'image')

    # Settings sub-areas
    sub = [
        (0.84, 0.82, 'Hospital Profile\nName · Address\nLogo · Reg. No.', PRIMARY),
        (0.84, 0.65, 'Department Mgmt\nAdd / Edit / Remove', TEAL),
        (0.84, 0.48, 'Doctor Profiles\nDept · Qualification', GREEN),
    ]
    ax.text(0.84, 0.93, 'Settings Sections', ha='center', fontsize=8, color=DARK, fontweight='bold')
    for (x, y, txt, bg) in sub:
        draw_box(ax, x, y, 0.18, 0.10, txt, bg=bg, fg=WHITE, fontsize=7)
    return fig_to_bytes(fig)

# ══════════════════════════════════════════════════════════════════════════════
# FLOWCHART 13 — Discharge Summary
# ══════════════════════════════════════════════════════════════════════════════
def chart_discharge_summary():
    fig, ax = new_fig(10, 4.5)
    add_title(ax, 'Discharge Summary — Template to PDF Workflow')

    steps = [
        (0.08, 0.60, 'Admin uploads\nDischarge Template\n(PDF / HTML)', GRAY),
        (0.24, 0.60, 'Template saved\nto library\nby category', PRIMARY),
        (0.40, 0.60, 'Doctor selects\ntemplate for\npatient', TEAL),
        (0.56, 0.60, 'DischargeEditor\nopens\n(HTML Editor)', PURPLE),
        (0.72, 0.60, 'Auto-fill:\nName · UHID · Age\nDoctor · Dept · Bed\nAdmission dates', BLUE),
        (0.88, 0.60, 'Doctor edits:\nDiagnosis\nMedications\nInstructions\nFollow-up', AMBER),
    ]
    for (x, y, txt, bg) in steps:
        draw_box(ax, x, y, 0.13, 0.24, txt, bg=bg, fg=WHITE, fontsize=7)

    for i in range(5):
        x1 = 0.08 + i*0.16 + 0.065
        arrow(ax, x1, 0.60, x1+0.03, 0.60)

    draw_box(ax, 0.55, 0.22, 0.18, 0.10, 'Save to server\nStatus: In Progress', TEAL, WHITE, 7.5)
    draw_box(ax, 0.78, 0.22, 0.18, 0.10, 'Download PDF\nhtml2pdf.js\nStatus: COMPLETED', GREEN, WHITE, 7.5)

    arrow(ax, 0.88, 0.48, 0.62, 0.27)
    arrow(ax, 0.64, 0.22, 0.69, 0.22)
    return fig_to_bytes(fig)


# ══════════════════════════════════════════════════════════════════════════════
# BUILD WORD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
def set_col_widths(table, widths):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = widths[i]

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    if level == 1:
        h.runs[0].font.color.rgb = RGBColor(0x6B, 0x21, 0xA8)  # deep purple
    elif level == 2:
        h.runs[0].font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)  # teal
    return h

def add_para(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 * (level + 1))
    return p

def add_flowchart(doc, img_bytes, caption='', width=Inches(6.2)):
    doc.add_picture(img_bytes, width=width)
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(8)
        cap.runs[0].italic = True
        cap.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '6366F1')
        tcPr.append(shd)
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    if col_widths:
        set_col_widths(table, col_widths)
    return table

def add_step_box(doc, steps_data):
    """steps_data: list of (number, title, description)"""
    table = doc.add_table(rows=1, cols=len(steps_data))
    table.style = 'Table Grid'
    row = table.rows[0].cells
    colors = ['6366F1', '0D9488', '8B5CF6', '10B981']
    for i, (num, title, desc) in enumerate(steps_data):
        c = row[i]
        tc = c._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), colors[i % len(colors)])
        tcPr.append(shd)
        p = c.paragraphs[0]
        p.clear()
        r = p.add_run(f'Step {num}\n{title}\n')
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(8)
        r2 = p.add_run(desc)
        r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r2.font.size = Pt(7)

# ══════════════════════════════════════════════════════════════════════════════
def build_document():
    print("[INFO] Generating flowchart images...")
    imgs = {
        'arch':      chart_system_arch(),
        'auth':      chart_auth(),
        'lifecycle': chart_patient_lifecycle(),
        'reg':       chart_registration_wizard(),
        'ipd':       chart_ipd(),
        'emergency': chart_emergency(),
        'lab':       chart_lab(),
        'forms':     chart_forms(),
        'billing':   chart_billing(),
        'cpoe':      chart_cpoe(),
        'appts':     chart_appointments(),
        'settings':  chart_settings(),
        'discharge': chart_discharge_summary(),
    }
    print("[OK] All flowcharts generated")

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width  = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ── Title Page ────────────────────────────────────────────────────────────
    title = doc.add_heading('Working Functionalities of SwasthyaSync HMS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x6B, 0x21, 0xA8)
    title.runs[0].font.size = Pt(22)

    doc.add_paragraph()
    meta = [
        ('Product Name', 'SwasthyaSync'),
        ('Version', '1.0.0'),
        ('Document Type', 'Working Functionalities Reference'),
        ('Technology Stack', 'React 18 (Vite) · Node.js (Express) · PostgreSQL'),
        ('Date', 'June 2026'),
    ]
    t = doc.add_table(rows=len(meta), cols=2)
    t.style = 'Table Grid'
    for i, (k, v) in enumerate(meta):
        row = t.rows[i].cells
        row[0].text = k
        row[0].paragraphs[0].runs[0].font.bold = True
        row[1].text = v
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 1 — System Overview
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '1. System Overview')
    add_para(doc, (
        'SwasthyaSync is a full-stack, cloud-native Hospital Management System (HMS) designed to '
        'digitise and unify all hospital workflows into a single integrated platform. '
        'The system covers 40+ hospital modules spanning patient registration, clinical care, '
        'diagnostics, pharmacy, billing, revenue cycle, and administration.'
    ))

    add_heading(doc, 'System Architecture', 2)
    add_flowchart(doc, imgs['arch'], 'Figure 1.1 — High-Level System Architecture')
    doc.add_paragraph()

    add_heading(doc, 'Key Highlights', 2)
    add_table(doc,
        ['Feature', 'Details'],
        [
            ['Modules', '40+ fully integrated hospital modules'],
            ['Compliance', 'NABH documentation standards, Insurance audit trails'],
            ['Access Control', 'Role-based access control (Admin / Staff)'],
            ['Security', 'JWT-based authentication, bcrypt-encrypted credentials'],
            ['PDF Generation', 'Built-in lab reports, discharge summaries, invoices'],
            ['Form System', 'PDF annotation engine for hospital consent forms'],
            ['Deployment', 'Docker-ready, Vite-built frontend, Express backend'],
        ]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 2 — Authentication
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, '2. Authentication & Access Control')
    add_flowchart(doc, imgs['auth'], 'Figure 2.1 — Authentication & Role-Based Access Control Flow')
    doc.add_paragraph()

    add_heading(doc, '2.1 Login', 2)
    for b in [
        'Email + password authentication',
        'JWT tokens issued on login and stored in browser localStorage',
        'Token verified on every protected route via AuthContext',
        'Automatic redirect to /login for unauthenticated sessions',
        'Loading state prevents flash of unauthenticated content',
    ]:
        add_bullet(doc, b)

    add_heading(doc, '2.2 Registration', 2)
    for b in [
        'New hospital/user registration form',
        'Fields: Hospital name, user name, email, password, confirm password',
        'Password hashing handled server-side (bcrypt)',
    ]:
        add_bullet(doc, b)

    add_heading(doc, '2.3 Role-Based Access Control', 2)
    add_table(doc,
        ['Role', 'Access Level', 'Restricted Modules'],
        [
            ['Admin', 'Full system access', 'None — access to all 40+ modules'],
            ['Staff', 'Operational access', 'Staff Management, Form Templates, Audit Log'],
        ]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 3 — Dashboard
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, '3. Dashboard & Analytics')
    add_para(doc, 'The main authenticated landing page provides hospital-wide at-a-glance metrics with live data from all connected modules.')

    add_heading(doc, 'KPI Stat Cards', 2)
    add_table(doc,
        ['Metric', 'Description', 'Trend'],
        [
            ['Total Patients', 'All patients registered in the system', '↑ % vs last period'],
            ['Bed Occupancy', 'Current % with bed detail (occupied/total)', '↑ % change'],
            ['Pending Notes', 'Clinical notes awaiting digitisation', '↓ (target: reduce)'],
            ['Today OPD', 'Outpatient visits today', '↑ % change'],
            ['Revenue (Month)', 'Monthly billed revenue vs. last month', '↑/↓ %'],
            ['Pending Lab Tests', 'Lab orders awaiting results', 'Count'],
            ['Critical Patients', 'Patients flagged as critical', 'Count'],
            ['Today Revenue', 'Billing amount generated today', 'Rs value'],
        ]
    )

    add_heading(doc, 'Module Launcher — 13 Category Tabs', 2)
    add_table(doc,
        ['Category Tab', 'Key Modules'],
        [
            ['Patient Administration', 'Registration, OPD, IPD, Queue, QR, Barcode, Merging, Medical Records'],
            ['Clinical', 'Doctors Workbench, Clinical Notes, Discharge Summary, Patient Forms'],
            ['Patient Billing', 'OPD Billing, IPD Billing, Payment Collection, GST Management'],
            ['Revenue Cycle', 'Claims Processing, Claim Submission, Medical Audit, Alerts'],
            ['Radiology', 'Radiology Orders, Report Management, PACS Integration'],
            ['Laboratory', 'Lab Orders, Result Entry, Sample Collection, Lab Reports'],
            ['Nursing Management', 'Nurse Station, Medication Admin, Vitals Monitoring, Ward Rounds'],
            ['Operation Theatre', 'OT Scheduling, OT Notes, Anaesthesia Notes, OT Reports'],
            ['Blood Bank', 'Blood Inventory, Donor Management, Blood Requests, Reports'],
            ['Pharmacy', 'Drug Dispensing, Drug Inventory, Pharmacy Billing, Prescriptions'],
            ['Inventory Management', 'Stock Management, Purchase Orders, Vendor Management'],
            ['Analytics', 'MIS Dashboard, MIS Reports, Revenue Analytics, Operational Reports'],
            ['Emergency', 'Emergency Triage, Ambulance Management, Critical Care'],
        ]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 4 — Patient Administration
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, '4. Patient Administration')

    add_heading(doc, '4.1 Patient Lifecycle', 2)
    add_flowchart(doc, imgs['lifecycle'], 'Figure 4.1 — Patient End-to-End Lifecycle')
    doc.add_paragraph()

    add_heading(doc, '4.2 Patient Registration — 4-Step Wizard', 2)
    add_flowchart(doc, imgs['reg'], 'Figure 4.2 — New Patient Registration Wizard Flow')
    doc.add_paragraph()

    add_step_box(doc, [
        ('1', 'Personal Details', 'Name · Age · Gender · Blood Group · Contact · Address · ABHA · Guardian'),
        ('2', 'Clinical Info', 'Department · Doctor · Status · MLC Type · Chief Complaint'),
        ('3', 'Payment/Insurance', 'Payment Type · Insurance · TPA · Policy No · Validity'),
        ('4', 'Consent & Submit', 'Summary review · Consent checkbox · Estimate provided · UHID issued'),
    ])
    doc.add_paragraph()

    add_heading(doc, 'Step 1 — Personal Details Fields', 2)
    add_table(doc,
        ['Field', 'Details', 'Required'],
        [
            ['Full Name', 'Patient full name', 'Yes'],
            ['Admission Type', 'OPD / IPD / Emergency / Day Care / ICU Direct', 'Yes'],
            ['Age', 'Numeric (years)', 'Yes'],
            ['Date of Birth', 'Date picker', 'No'],
            ['Gender', 'Male / Female / Other / Prefer not to say', 'Yes'],
            ['Blood Group', 'A+/A-/B+/B-/AB+/AB-/O+/O-/Unknown', 'No'],
            ['Mobile Number', '10-digit mobile', 'Yes'],
            ['Aadhaar Number', 'XXXX XXXX XXXX format', 'No'],
            ['ABHA Health ID', '14-digit ABHA number', 'No'],
            ['Address', 'Street, city, state (all Indian states), PIN', 'No'],
            ['Guardian Name / Relation / Phone', 'Emergency contact details', 'No'],
        ]
    )

    add_heading(doc, 'Step 2 — Clinical Info Fields', 2)
    add_table(doc,
        ['Field', 'Details'],
        [
            ['Department', '30+ departments (General Medicine, ICU, Cardiology, Neurology, etc.)'],
            ['Attending Doctor', 'Dropdown from registered doctors'],
            ['Initial Status', 'Stable / Critical / Recovering / Under Obs / Serious'],
            ['Patient Category', 'General / BPL / Senior Citizen / Divyangjan / VIP / Staff'],
            ['MLC Type', 'None / Road Accident / Assault / Poisoning / Burns / Suicide Attempt / etc.'],
            ['Chief Complaint', 'Free-text textarea for presenting symptoms'],
        ]
    )

    add_heading(doc, 'Step 3 — Payment / Insurance Fields', 2)
    add_table(doc,
        ['Field', 'Details'],
        [
            ['Payment Type', 'Self Pay (Cash/UPI) / Insurance/TPA / CGHS / ECHS / ESI / Ayushman Bharat / Govt/Free'],
            ['Insurance Company', 'Visible only when Insurance/TPA selected'],
            ['TPA Name', 'Third-party administrator name'],
            ['Policy / Member ID', 'Policy number'],
            ['Policy Validity', 'Expiry date picker'],
        ]
    )

    add_heading(doc, '4.3 Patient Detail Panel — Tabs', 2)
    add_table(doc,
        ['Tab', 'Contents'],
        [
            ['Overview', 'MLC alert banner, key info grid, chief complaint, quick navigation to Forms tab'],
            ['Clinical', 'Full demographics, Aadhaar, ABHA ID, insurance details section'],
            ['Forms', 'Assigned form instances, assign new template, open FormViewer PDF annotator'],
            ['Discharge', 'Assigned discharge summaries, assign template, open DischargeEditor HTML editor'],
        ]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 5 — IPD
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, '5. Inpatient Department (IPD)')
    add_flowchart(doc, imgs['ipd'], 'Figure 5.1 — IPD Bed Management & Discharge Workflow')
    doc.add_paragraph()

    add_heading(doc, 'Bed Status Definitions', 2)
    add_table(doc,
        ['Status', 'Colour', 'Description', 'Action on Click'],
        [
            ['Occupied', '🟣 Purple', 'Patient assigned: shows name, age, doctor, diagnosis', 'Open Bed Detail Panel'],
            ['Available', '🟢 Green', 'No patient assigned', 'Open Assign Bed modal'],
            ['Alert', '🔴 Red dot', 'Pulsing indicator — critical patient', 'Open Bed Detail Panel'],
            ['Maintenance', '⚙️ Gray', 'Under repair/cleaning', 'No action'],
            ['Reserved', '🟡 Amber', 'Reserved for upcoming admission', 'No action'],
        ]
    )

    add_heading(doc, 'Bed Detail Panel Tabs', 2)
    add_table(doc,
        ['Tab', 'Contents & Actions'],
        [
            ['Patient Info', 'Ward, bed type, attending doctor, admission date, diagnosis. Quick links to Forms and Discharge tabs. Buttons: View Full Record, Write Clinical Note, Generate Discharge Summary, Release Bed'],
            ['Forms', 'List of assigned forms with status (Blank/In Progress/Completed). Assign new templates from categorised library. Opens FormViewer on click.'],
            ['Discharge', 'Assign discharge summary templates. Opens DischargeEditor (HTML rich text). Download as PDF. Track status.'],
        ]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 6 — Emergency
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, '6. Emergency Module')
    add_flowchart(doc, imgs['emergency'], 'Figure 6.1 — Emergency Triage & MLC Workflow')
    doc.add_paragraph()

    add_heading(doc, 'Triage Levels', 2)
    add_table(doc,
        ['Level', 'Colour', 'Description', 'Response Time'],
        [
            ['Red — Critical', '🔴', 'Life-threatening. Alert flag ON. Immediate care required.', 'Immediate'],
            ['Orange — Urgent', '🟠', 'Serious condition. Priority care.', 'Within 15 minutes'],
            ['Yellow — Semi-Urgent', '🟡', 'Moderate condition. Can wait briefly.', 'Within 60 minutes'],
            ['Green — Non-Urgent', '🟢', 'Minor condition. Standard OPD queue.', 'Standard queue'],
        ]
    )

    add_heading(doc, 'MLC (Medico-Legal Case) Types', 2)
    add_table(doc,
        ['MLC Type'],
        [
            ['Road Traffic Accident (RTA)'],
            ['Assault'],
            ['Poisoning'],
            ['Burns'],
            ['Sexual Assault'],
            ['Suicide Attempt'],
            ['Industrial Accident'],
            ['Other MLC'],
        ]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 7 — Appointments
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, '7. Appointments & Scheduling')
    add_flowchart(doc, imgs['appts'], 'Figure 7.1 — Appointment Booking Workflow')
    doc.add_paragraph()

    add_table(doc,
        ['Status', 'Description'],
        [
            ['Scheduled', 'Appointment booked, patient not yet arrived'],
            ['Confirmed', 'Patient has arrived and checked in'],
            ['Completed', 'Consultation done, OPD workflow begins'],
            ['Cancelled', 'Patient did not arrive or appointment was cancelled'],
        ]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 8 — Clinical Modules
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, '8. Clinical Modules')

    add_heading(doc, '8.1 CPOE — Computerised Physician Order Entry', 2)
    add_flowchart(doc, imgs['cpoe'], 'Figure 8.1 — CPOE Order Entry Flow')
    doc.add_paragraph()

    add_heading(doc, '8.2 Discharge Summary Workflow', 2)
    add_flowchart(doc, imgs['discharge'], 'Figure 8.2 — Discharge Summary Template to PDF Workflow')
    doc.add_paragraph()

    add_heading(doc, '8.3 Patient Forms & Consent Management — FormViewer', 2)
    add_flowchart(doc, imgs['forms'], 'Figure 8.3 — FormViewer PDF Annotation Engine Workflow')
    doc.add_paragraph()

    add_heading(doc, 'Form Categories', 2)
    add_table(doc,
        ['Category', 'Typical Forms'],
        [
            ['Consent', 'General Consent for Treatment, Surgical Consent, Anaesthesia Consent'],
            ['Assessment', 'Initial Assessment Form, Nursing Assessment, Pain Assessment'],
            ['Nursing', 'Nursing Notes, Medication Administration Record (MAR)'],
            ['ICU', 'ICU Admission Form, Critical Care Daily Assessment'],
            ['OT', 'Pre-operative Checklist, OT Consent, Post-op Record'],
            ['Emergency', 'Emergency Assessment, MLC Documentation Form'],
            ['Discharge', 'Discharge Instructions, Against Medical Advice (AMA) Form'],
            ['General', 'Patient Rights & Responsibilities, Hospital Policy Acknowledgement'],
        ]
    )

    add_heading(doc, 'Clinical Note Types', 2)
    add_table(doc,
        ['Note Type', 'Description'],
        [
            ['Admission Note', 'Initial assessment on patient admission'],
            ['Progress Note', 'Daily/shift update on patient condition'],
            ['Procedure Note', 'Documentation of a clinical procedure performed'],
            ['Consultant Note', 'Opinion from a specialist consultant'],
            ['Nursing Note', 'Nursing observations and care provided'],
        ]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 9 — Laboratory
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, '9. Laboratory Information System (LIS)')
    add_flowchart(doc, imgs['lab'], 'Figure 9.1 — Lab Order to PDF Report Complete Workflow')
    doc.add_paragraph()

    add_heading(doc, 'Lab Test Catalog — 50+ Tests by Category', 2)
    add_table(doc,
        ['Category', 'Tests (examples)'],
        [
            ['Haematology', 'CBC, ESR, PT/INR, APTT, D-Dimer, G6PD, Blood Group & Rh'],
            ['Biochemistry', 'LFT, KFT, Lipid Profile, Blood Sugar (Fasting/PP/Random), ABG, Amylase, Lipase, Electrolytes'],
            ['Pathology', 'Thyroid Profile (T3/T4/TSH), HbA1c, Vitamin D, B12, PSA, CEA, CA125, CA19-9, AFP, Beta HCG'],
            ['Microbiology', 'Widal, Dengue NS1/IgG/IgM, Malaria Antigen, HIV, HBsAg, HCV, VDRL, Blood/Urine Cultures, Mantoux'],
            ['Immunology', 'CRP, RA Factor, ASO Titre'],
            ['Cardiology', 'Troponin I/T, CPK, CPK-MB, LDH'],
        ]
    )

    add_heading(doc, 'Lab Order Priority Levels', 2)
    add_table(doc,
        ['Priority', 'Description', 'Turnaround'],
        [
            ['STAT', 'Life-threatening emergency — process immediately', 'Immediate'],
            ['URGENT', 'Clinically significant — process as soon as possible', 'Within 1-2 hours'],
            ['ROUTINE', 'Standard testing — normal processing queue', 'Standard TAT'],
        ]
    )

    add_heading(doc, 'Lab Report PDF Contents', 2)
    add_table(doc,
        ['PDF Section', 'Contents'],
        [
            ['Header', 'Hospital logo + name + address + tagline (Text/Logo mode) OR full header image (Image mode)'],
            ['Report Title', 'LABORATORY TEST REPORT'],
            ['Patient Info Box', 'Patient Name, Patient ID, Requesting Doctor, Report Date & Time'],
            ['Investigation Title', 'Test name (e.g. CBC, LFT, etc.)'],
            ['Results Table', 'Parameter | Result | Flag (HIGH/LOW) | Unit | Biological Reference Range'],
            ['Flag Indicators', 'HIGH values in Red, LOW values in Amber, Normal values in Black'],
            ['Footer', 'Lab director signature, accreditation info, page numbers (Text mode) OR full footer image'],
        ]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 10 — Billing
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, '10. Billing & Revenue Cycle')
    add_flowchart(doc, imgs['billing'], 'Figure 10.1 — Complete Billing & Revenue Cycle Workflow')
    doc.add_paragraph()

    add_heading(doc, 'Invoice Status Types', 2)
    add_table(doc,
        ['Status', 'Description'],
        [
            ['Paid ✅', 'Full payment collected. Bill settled.'],
            ['Pending ⏳', 'Invoice generated, payment not yet received.'],
            ['Partial 💳', 'Partial payment received, balance outstanding.'],
            ['Overdue ⚠', 'Payment due date passed, not yet settled.'],
        ]
    )

    add_heading(doc, 'Accepted Payment Methods', 2)
    add_table(doc,
        ['Method', 'Description'],
        [
            ['Cash', 'Cash payment at counter'],
            ['Card', 'Debit/Credit card payment'],
            ['UPI', 'UPI payment (GPay, PhonePe, Paytm etc.)'],
            ['Insurance', 'Third-party insurance settlement'],
            ['Cheque', 'Cheque payment'],
            ['NEFT', 'Bank transfer via NEFT/RTGS'],
        ]
    )

    add_heading(doc, 'IPD Discharge Billing — 6-Step Clearance Workflow', 2)
    add_table(doc,
        ['Step', 'Stage', 'Description'],
        [
            ['0', 'Discharge Advised', 'Doctor advises discharge. Process initiated.'],
            ['1', 'Clinical Summary', 'Doctor completes clinical summary and signs off on case.'],
            ['2', 'Pharmacy Clearance', 'All pharmacy charges verified and added to bill.'],
            ['3', 'Laboratory Clearance', 'All lab test charges verified and added to bill.'],
            ['4', 'Final Billing', 'All charges consolidated: ward, nursing, OT, lab, pharmacy, radiology.'],
            ['5', 'Settle & Vacate', 'Patient pays final bill. Bed released. Admission closed.'],
        ]
    )

    add_heading(doc, 'TPA / Insurance Management', 2)
    for b in [
        'Insurance company and TPA master data management',
        'Pre-authorisation request tracking and submission',
        'Claims processing and submission to insurance company',
        'Medical audit trail — full documentation for insurance review',
        'Technical audit documentation for claim processing',
        'Alerts management for claim status updates and queries',
    ]:
        add_bullet(doc, b)

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 11 — Settings
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, '11. Settings & Hospital Configuration')
    add_flowchart(doc, imgs['settings'], 'Figure 11.1 — Settings Module & Report Branding Decision Flow')
    doc.add_paragraph()

    add_heading(doc, 'Report Print Modes', 2)
    add_table(doc,
        ['Mode', 'Header', 'Footer', 'Use Case'],
        [
            ['Text / Logo', 'Hospital name + tagline + logo from settings', 'Text footer + page numbers', 'Standard branded reports'],
            ['Image Mode', 'Full-page header image (uploaded by admin)', 'Full-page footer image (uploaded by admin)', 'Pre-designed letterhead/stationery'],
        ]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 12 — Other Modules
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, '12. Other Clinical & Operational Modules')

    add_heading(doc, '12.1 Radiology Information System (RIS)', 2)
    add_table(doc,
        ['Feature', 'Details'],
        [
            ['Order Creation', 'Radiology order linked to patient and requesting doctor'],
            ['Investigation Types', 'X-Ray, CT Scan, MRI, Ultrasound, ECG, Echo, and more'],
            ['PACS Integration', 'Integration point for PACS (Picture Archiving and Communication System)'],
            ['Report Management', 'Type report in system or upload scanned PDF'],
            ['Status Tracking', 'Pending → Scanned → Reported'],
        ]
    )

    add_heading(doc, '12.2 Pharmacy Management', 2)
    add_table(doc,
        ['Feature', 'Details'],
        [
            ['Drug Dispensing', 'Prescription-based drug dispensing from doctor orders'],
            ['Drug Inventory', 'Stock tracking with low-stock alerts'],
            ['Pharmacy Billing', 'Drug charges automatically added to patient bill'],
            ['Drug Catalog', 'Medicine master with dosage and category'],
        ]
    )

    add_heading(doc, '12.3 Blood Bank', 2)
    add_table(doc,
        ['Feature', 'Details'],
        [
            ['Inventory Management', 'Blood stock tracking by blood group (A+/A-/B+/B-/AB+/AB-/O+/O-)'],
            ['Donor Management', 'Donor registration, donation history, eligibility check'],
            ['Blood Components', 'Whole Blood, Packed RBCs, Platelets, FFP, Cryoprecipitate'],
            ['Blood Requests', 'Request linked to patient, compatibility verification'],
        ]
    )

    add_heading(doc, '12.4 Operation Theatre (OT)', 2)
    add_table(doc,
        ['Stage', 'Documentation'],
        [
            ['Scheduling', 'OT booking: date, time, OT room, surgeon, anaesthesiologist'],
            ['Pre-operative', 'Surgery checklist, pre-op notes, patient briefing'],
            ['Anaesthesia', 'Anaesthesia notes by anaesthesiologist'],
            ['Intra-operative', 'OT notes: surgical findings, procedure steps'],
            ['Post-operative', 'Recovery room monitoring, post-op instructions'],
        ]
    )

    add_heading(doc, '12.5 Inventory Management', 2)
    for b in [
        'Stock item master: medicines, consumables, medical equipment',
        'Purchase order creation and vendor tracking',
        'Goods receipt and automatic stock increment on receipt',
        'Stock level monitoring with configurable low-stock alerts',
        'Vendor master management with contact and payment terms',
        'Reports: current stock, stock movement, inventory valuation',
    ]:
        add_bullet(doc, b)

    add_heading(doc, '12.6 Staff Management (Admin Only)', 2)
    for b in [
        'Staff registration: doctors, nurses, administrative staff, lab technicians',
        'Role assignment: Admin or Staff',
        'Department allocation and designation management',
        'Active / Inactive status control',
        'Staff list with search and filter',
    ]:
        add_bullet(doc, b)

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 13 — Reports & MIS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '13. Reports & MIS', 1)
    add_table(doc,
        ['Report Type', 'Contents'],
        [
            ['MIS Dashboard', 'Hospital-level KPI overview with live metrics'],
            ['MIS Reports', 'Exportable data: admissions, discharges, OPD visits'],
            ['Revenue Analytics', 'Revenue trend, department-wise collections, insurance vs. cash split'],
            ['Operational Reports', 'Bed occupancy, average length of stay, OT utilisation, lab TAT'],
            ['Clinical Audit', 'Quality metrics, clinical outcome tracking, NABH documentation'],
        ]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 14 — Audit Log
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '14. Audit Log (Admin Only)', 1)
    add_para(doc, 'Complete chronological log of all system actions, tamper-proof for compliance and investigation purposes.')
    add_table(doc,
        ['Event Type', 'Details Logged'],
        [
            ['Patient created / updated', 'UHID, fields changed, user, timestamp'],
            ['Bed assigned / released', 'Bed ID, patient, doctor, action, timestamp'],
            ['Lab order / result entered', 'Order ID, test name, result, user, timestamp'],
            ['Bill generated / payment received', 'Invoice ID, amount, payment method, user'],
            ['Form assigned / completed', 'Form ID, patient, template, status change'],
            ['User login / logout', 'Username, IP, timestamp'],
            ['Settings changed', 'Field changed, old value, new value, user'],
        ]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 15 — Technical Architecture
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, '15. Technical Architecture')

    add_heading(doc, 'Frontend Technology Stack', 2)
    add_table(doc,
        ['Component', 'Technology', 'Purpose'],
        [
            ['Framework', 'React 18 with Vite', 'UI framework with fast HMR development'],
            ['Routing', 'React Router v6', 'Client-side routing with protected routes'],
            ['Charts', 'Recharts', 'Area, Bar, and Pie charts for analytics'],
            ['Icons', 'Lucide React', 'Consistent icon library'],
            ['PDF Generation', 'jsPDF + jsPDF-AutoTable + html2pdf.js', 'Lab reports, discharge summaries, invoices'],
            ['PDF Rendering', 'pdfjs-dist', 'Render PDF forms in browser'],
            ['QR Code', 'html5-qrcode', 'QR code scanning for patient registration'],
            ['State Management', 'React Context (AuthContext)', 'Global auth state management'],
            ['Styling', 'Vanilla CSS + CSS custom properties', 'Design system with tokens'],
        ]
    )

    add_heading(doc, 'Backend Technology Stack', 2)
    add_table(doc,
        ['Component', 'Technology', 'Purpose'],
        [
            ['Runtime', 'Node.js', 'JavaScript server runtime'],
            ['Framework', 'Express.js', 'REST API server'],
            ['Database', 'PostgreSQL', 'Relational data storage'],
            ['Authentication', 'JWT (jsonwebtoken) + bcrypt', 'Token auth + password hashing'],
            ['File Storage', 'Local filesystem /server/uploads/', 'PDFs, images, logos'],
            ['Email', 'Nodemailer (SMTP)', 'Lab report email delivery'],
        ]
    )

    add_heading(doc, 'Key Database Tables', 2)
    add_table(doc,
        ['Table', 'Purpose'],
        [
            ['patients', 'Master patient record with all demographics, clinical info, payment details'],
            ['beds', 'IPD bed registry with status, ward, patient assignment'],
            ['lab_orders', 'Lab test orders, results (JSON), PDF path, status'],
            ['radiology_orders', 'Radiology investigation orders and reports'],
            ['billing', 'Invoice and payment records with status tracking'],
            ['appointments', 'OPD appointment bookings with slot management'],
            ['doctors', 'Doctor master: name, department, qualification'],
            ['clinical_notes', 'Doctor clinical documentation with priority and type'],
            ['form_templates', 'Uploaded PDF form templates with category'],
            ['patient_forms', 'Patient-specific form instances with annotations JSON'],
            ['discharge_templates', 'Discharge summary template library'],
            ['discharge_summaries', 'Patient-specific discharge summary instances'],
            ['staff', 'Staff and user records with role and department'],
            ['audit_log', 'Tamper-proof system action audit trail'],
            ['hospital_settings', 'Hospital profile, branding, report configuration'],
        ]
    )

    add_heading(doc, 'Deployment Configuration', 2)
    add_table(doc,
        ['Component', 'Details'],
        [
            ['Containerisation', 'Docker + Docker Compose'],
            ['Frontend build', 'npm run build (Vite production bundle)'],
            ['Dev server', 'npm run dev (Vite with HMR on port 5173)'],
            ['Backend', 'Node.js Express server (port configurable via .env)'],
            ['Database', 'PostgreSQL (Docker container or managed instance)'],
            ['Environment', '.env file for DB credentials, JWT secret, SMTP config'],
        ]
    )

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run('Document prepared from codebase analysis of the SwasthyaSync HMS project.\n© 2026 SwasthyaSync. All rights reserved.')
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    out_path = r'c:\Users\suman\OneDrive\Desktop\DSCRIBE CLONE\Working_Functionalities_of_HMS.docx'
    doc.save(out_path)
    print(f'\n[DONE] Document saved to:\n   {out_path}')
    return out_path

if __name__ == '__main__':
    build_document()
